from __future__ import annotations

import json
import sqlite3
import tempfile
import unittest
from datetime import datetime
from pathlib import Path

import yaml
from docx import Document

from sagp_admin_processor import AdminProcessor, JobKind
from sagp_admin_processor.workflows import AppliedChange


ROOT = Path(__file__).resolve().parents[2]
SCHEMA = ROOT / "sagp_member_import/schema/sagp_members_schema.sql"


class WorkflowTests(unittest.TestCase):
    def setUp(self):
        self.tempdir = tempfile.TemporaryDirectory()
        self.root = Path(self.tempdir.name)
        (self.root / "sagp_member_manager/output").mkdir(parents=True)
        (self.root / "sagp_website/public/platform").mkdir(parents=True)
        (self.root / "sagp_website/content/events/annual_conference").mkdir(parents=True)
        (self.root / "sagp_website/content/calls").mkdir(parents=True)
        self.db = self.root / "sagp_member_manager/output/sagp_members.db"
        with sqlite3.connect(self.db) as con:
            con.executescript(SCHEMA.read_text(encoding="utf-8"))
        self.processor = AdminProcessor(self.root, self.root / "state", run_artifact_builders=False)

    def tearDown(self):
        self.tempdir.cleanup()

    def write_json(self, name: str, data: dict) -> Path:
        path = self.root / name
        path.write_text(json.dumps(data), encoding="utf-8")
        return path

    def insert_member(self):
        now = datetime.now().isoformat()
        with sqlite3.connect(self.db) as con:
            con.execute(
                "INSERT INTO members (person_id, display_name, first_name, last_name, primary_email, institution, last_paid_year, membership_status, active, created_at, updated_at) VALUES (?, ?, ?, ?, ?, ?, ?, ?, 1, ?, ?)",
                ("SAGP000010", "Ada Lovelace", "Ada", "Lovelace", "ada@example.org", "Analytical Society", 2025, "Current Member", now, now),
            )

    def test_membership_update_preview_process_and_rollback(self):
        self.insert_member()
        request = self.write_json("membership_update.json", {
            "person_id": "SAGP000010",
            "proposed_changes": {"institution": "Royal Society"},
            "reason": "Correction",
            "requested_by": "secretary",
            "request_id": "mur_test123456",
            "requested_at": datetime.now().astimezone().isoformat(),
            "source": "test",
            "notes": None,
        })
        preview = self.processor.preview_file(request)
        self.assertEqual(preview.kind, JobKind.MEMBERSHIP_UPDATE)
        self.assertTrue(preview.valid)
        self.processor.process(request, preview)
        with sqlite3.connect(self.db) as con:
            self.assertEqual(con.execute("SELECT institution FROM members WHERE person_id='SAGP000010'").fetchone()[0], "Royal Society")
        self.processor.back_out_last()
        with sqlite3.connect(self.db) as con:
            self.assertEqual(con.execute("SELECT institution FROM members WHERE person_id='SAGP000010'").fetchone()[0], "Analytical Society")

    def test_membership_create_detects_duplicate(self):
        self.insert_member()
        request = self.write_json("membership_create.json", {
            "request_type": "membership_create",
            "proposed_member_data": {"first_name": "Different", "last_name": "Person", "primary_email": "ADA@example.org", "institution": "Elsewhere", "region": "", "notes": "", "membership_expiration_year": "2027"},
            "reason": "Joined",
            "requested_by": "secretary",
            "request_id": "mcr_test123456",
            "requested_at": datetime.now().astimezone().isoformat(),
            "source": "test",
            "notes": None,
        })
        preview = self.processor.preview_file(request)
        self.assertEqual(preview.kind, JobKind.MEMBERSHIP_CREATE)
        self.assertFalse(preview.valid)
        self.assertTrue(any("email" in error.lower() for error in preview.errors))

    def test_membership_create_process_and_rollback(self):
        request = self.write_json("membership_create.json", {
            "request_type": "membership_create",
            "proposed_member_data": {"first_name": "Grace", "last_name": "Hopper", "primary_email": "grace@example.org", "institution": "US Navy", "region": "", "notes": "", "membership_expiration_year": "2027"},
            "reason": "Joined",
            "requested_by": "secretary",
            "request_id": "mcr_test654321",
            "requested_at": datetime.now().astimezone().isoformat(),
            "source": "test",
            "notes": None,
        })
        preview = self.processor.preview_file(request)
        self.assertTrue(preview.valid)
        change = self.processor.process(request, preview)
        self.assertEqual(change.details["person_id"], "SAGP000001")
        with sqlite3.connect(self.db) as con:
            self.assertEqual(con.execute("SELECT COUNT(*) FROM members").fetchone()[0], 1)
        self.processor.back_out_last()
        with sqlite3.connect(self.db) as con:
            self.assertEqual(con.execute("SELECT COUNT(*) FROM members").fetchone()[0], 0)

    def test_publication_update_process_and_rollback(self):
        yaml_path = self.root / "sagp_website/content/events/annual_conference/2027.yaml"
        yaml_path.write_text(yaml.safe_dump({"event": {"id": "annual-conference-2027", "title": "Old Title", "status": "upcoming"}}, sort_keys=False), encoding="utf-8")
        request = self.write_json("publication_update.json", {
            "request_type": "publication_update",
            "request_id": "pur_test",
            "publication_type": "event",
            "publication_id": "annual-conference-2027",
            "publication_title": "Old Title",
            "reason": "Correct title",
            "context": {"source_yaml": "content/events/annual_conference/2027.yaml"},
            "changes": {"title": {"before": "Old Title", "after": "New Title"}},
        })
        preview = self.processor.preview_file(request)
        self.assertTrue(preview.valid)
        self.processor.process(request, preview)
        self.assertEqual(yaml.safe_load(yaml_path.read_text())["event"]["title"], "New Title")
        self.processor.back_out_last()
        self.assertEqual(yaml.safe_load(yaml_path.read_text())["event"]["title"], "Old Title")

    def test_publication_safety_mismatch_is_invalid(self):
        yaml_path = self.root / "sagp_website/content/events/annual_conference/2027.yaml"
        yaml_path.write_text(yaml.safe_dump({"event": {"title": "Current"}}), encoding="utf-8")
        request = self.write_json("publication_update.json", {
            "request_type": "publication_update", "publication_type": "event",
            "context": {"source_yaml": "content/events/annual_conference/2027.yaml"},
            "changes": {"title": {"before": "Stale", "after": "New"}},
        })
        preview = self.processor.preview_file(request)
        self.assertFalse(preview.valid)
        self.assertTrue(any("Safety check" in error for error in preview.errors))

    def test_publication_change_requires_before_and_after(self):
        yaml_path = self.root / "sagp_website/content/events/annual_conference/2027.yaml"
        yaml_path.write_text(yaml.safe_dump({"event": {"title": "Current"}}), encoding="utf-8")
        request = self.write_json("publication_update.json", {
            "request_type": "publication_update", "publication_type": "event",
            "context": {"source_yaml": "content/events/annual_conference/2027.yaml"},
            "changes": {"title": "New"},
        })
        preview = self.processor.preview_file(request)
        self.assertFalse(preview.valid)
        self.assertTrue(any("before and after" in error for error in preview.errors))

    def test_docx_requires_options_and_previews_content(self):
        path = self.root / "event.docx"
        document = Document()
        document.add_paragraph("2027 SAGP Annual Conference")
        document.add_paragraph("A conference about Plato")
        document.save(path)
        missing = self.processor.preview_file(path)
        self.assertFalse(missing.valid)
        preview = self.processor.preview_file(path, "annual_conference", "2027")
        self.assertTrue(preview.valid)
        self.assertEqual(preview.kind, JobKind.EVENT_DOCX)
        self.assertIn("A conference about Plato", preview.summary)

    def test_communication_json_is_explicitly_unsupported(self):
        path = self.write_json("communication.json", {"communication_id": "com_123"})
        preview = self.processor.preview_file(path)
        self.assertEqual(preview.kind, JobKind.UNSUPPORTED)
        self.assertIn("Communication", preview.summary)

    def test_deployment_orchestration_uses_separate_repositories(self):
        class RecordingProcessor(AdminProcessor):
            def __init__(self, *args, **kwargs):
                super().__init__(*args, **kwargs)
                self.commands = []
                self.website_staged = False
                self.root_staged = False

            def _run(self, cmd, cwd, log, allow_fail=False):
                self.commands.append((tuple(cmd), cwd))
                if cmd[:2] == ["git", "add"] and cwd == self.website:
                    self.website_staged = True
                elif cmd[:2] == ["git", "commit"] and cwd == self.website:
                    self.website_staged = False
                elif cmd[:2] == ["git", "add"] and cwd == self.root:
                    self.root_staged = True
                elif cmd[:2] == ["git", "commit"] and cwd == self.root:
                    self.root_staged = False

            def _output(self, cmd, cwd):
                if cmd[:3] == ["git", "diff", "--cached"]:
                    if cwd == self.website and self.website_staged:
                        return "public/platform/membership_records.json"
                    if cwd == self.root and self.root_staged:
                        return "sagp_website"
                return ""

        processor = RecordingProcessor(self.root, self.root / "deploy-state", run_artifact_builders=False)
        snapshot = self.root / "deploy-state/change_test"
        snapshot.mkdir(parents=True)
        processor.applied.append(AppliedChange(
            "change_test", JobKind.MEMBERSHIP_UPDATE, "Ada Lovelace",
            self.root / "request.json", snapshot, [],
            {"public/platform/membership_records.json"}, {},
        ))
        processor.deploy(watch=False)
        commands = [command for command, _ in processor.commands]
        self.assertIn(("npm", "run", "build"), commands)
        self.assertIn(("git", "push"), commands)
        self.assertEqual(commands.count(("git", "push")), 2)
        self.assertFalse(processor.applied)
        self.assertEqual(processor.deployment_phase, "local")


if __name__ == "__main__":
    unittest.main()
